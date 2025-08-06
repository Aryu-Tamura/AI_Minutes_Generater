import streamlit as st
from openai import OpenAI
import whisper
import torch
from pyannote.audio import Pipeline
from pydub import AudioSegment
import tempfile
import os
from datetime import timedelta, date, datetime
from docx import Document
from docx.shared import Inches, Pt
from io import BytesIO
import json
import plotly.graph_objects as go
import logging
import sqlite3
import zipfile
import re

# -------------------------------------------------------------------
# 1. 初期設定 & ロギング・DB設定
# -------------------------------------------------------------------

st.set_page_config(layout="wide", page_title="AI交渉アシスタント")

# ロギングの設定
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s', filename='app.log', filemode='a')

# StreamlitのsecretsからAPIキーを安全に読み込む
try:
    HF_TOKEN = st.secrets["HF_TOKEN"]
    OPENAI_API_KEY = st.secrets["OPENAI_API_KEY"]
    logging.info("API keys loaded successfully.")
except FileNotFoundError:
    st.error("`.streamlit/secrets.toml`ファイルが見つかりません。APIキーを設定してください。")
    logging.error("secrets.toml not found.")
    st.stop()
except KeyError as e:
    st.error(f"`secrets.toml`に`{e}`が設定されていません。")
    logging.error(f"API key missing in secrets.toml: {e}")
    st.stop()

# 各種クライアントの初期化
WHISPER_MODEL = "small"
client = OpenAI(api_key=OPENAI_API_KEY)

# データベースの初期化
DB_FILE = "database.db"
def init_db():
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    c.execute('''
        CREATE TABLE IF NOT EXISTS reports (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            timestamp TEXT NOT NULL,
            sales_rep TEXT NOT NULL,
            client_company TEXT NOT NULL,
            client_rep TEXT NOT NULL,
            report_date TEXT NOT NULL,
            analysis_json TEXT NOT NULL,
            report_markdown TEXT,
            cleaned_transcript TEXT
        )
    ''')
    conn.commit()
    conn.close()
    logging.info("Database initialized.")
init_db()

# -------------------------------------------------------------------
# 2. セッションステートの管理
# -------------------------------------------------------------------
def reset_creation_page_state():
    """商談レポート作成ページの状態をリセットする関数"""
    st.session_state.analysis_stage = "initial"
    st.session_state.negotiation_info = {}
    st.session_state.analysis_data = None
    st.session_state.transcript_display = []
    st.session_state.chat_history = []
    st.session_state.report_for_display = ""
    st.session_state.uploaded_file = None
    st.session_state.current_report_id = None
    st.session_state.report_saved = False
    logging.info("Creation page state has been reset.")

if "current_page" not in st.session_state:
    st.session_state.current_page = "creation"
    reset_creation_page_state()

# -------------------------------------------------------------------
# 3. OpenAI GPT API 関連関数
# -------------------------------------------------------------------

def get_negotiation_analysis(transcript_text, negotiation_info):
    """
    文字起こしデータから、根拠とセールスフローに基づいた交渉分析を行う関数。
    """
    system_prompt = """
あなたは、銀行渉外担当者のための超一流ネゴシエーション・コーチです。
提供された商談の文字起こしを、以下の**理想的なセールスフロー**と**採点基準（ルーブリック）**に照らし合わせて、厳格に評価してください。

### 理想的なセールスフロー
1.  **関係構築 (Rapport Building)**
2.  **課題発見 (Problem Discovery)**
3.  **価値提案 (Value Proposition)**
4.  **合意形成とクロージング (Closing)**

### 採点基準（スコアリング・ルーブリック）
-   **A (Excellent)**: 相手への共感や承認の言葉が豊かで、オープンな質問を通じて相手が本音を話しやすい雰囲気を作れている。
-   **B (Good)**: 丁寧な挨拶や共感の言葉は見られるが、会話を広げるための工夫がやや不足している。
-   **C (Average)**: 事務的なやり取りに終始し、相手の感情に寄り添う姿勢が見られない。
-   **D (Needs Improvement)**: 一方的な発言や、相手を否定するような言動が見られ、関係構築の機会を逃している。

あなたの任務は、会話の文脈全体を考慮し、安易な高評価を避け、**この採点基準に厳密に従ってA〜Dの評価を下す**ことです。

**【根拠引用の絶対ルール】**
根拠となった会話は、**必ず複数人の発言を含む「会話のキャッチボール」**を引用してください。単一の発言だけを引用することは許可されません。文脈を理解する上で十分な長さのやり取りを抜き出してください。
"""
    user_prompt = f"""
### 指示
上記の評価基準に基づき、以下の商談の文字起こしデータを分析し、各ステージの評価（A〜D）と分析内容をJSON形式で出力してください。

### 話者名の特定
営業担当者は「{negotiation_info['sales_rep']}」です。文字起こしデータ内の「SPEAKER_00」「SPEAKER_01」などを分析し、どちらが営業担当者でどちらが顧客（{negotiation_info['client_rep']}）かを判断してください。その上で、"cleaned_transcript"内の"speaker"を、実際の名前（例：「田中真奈美（営業担当）」、「藤社長」）に置き換えてください。

### 分析対象の文字起こしデータ
```
{transcript_text}
```

### 出力フォーマット (JSON)
```json
{{
  "cleaned_transcript": [
    {{ "speaker": "（話者名）", "text": "（発言内容）", "start_time": "（開始時間）" }}
  ],
  "summary_report": {{
    "overview": {{
        "date": "{negotiation_info['date']}",
        "attendees": {{
            "client_company": "{negotiation_info['client_company']}",
            "client_rep": "{negotiation_info['client_rep']} 様",
            "our_company": "{negotiation_info['sales_rep']}"
        }}
    }},
    "agenda": "（本日のアジェンダを要約）",
    "summary": [
        "（議論全体の要点を具体的に要約した1つ目の箇条書き）",
        "（議論全体の要点を具体的に要約した2つ目の箇条書き）"
    ],
    "decisions": ["（決定事項1）"],
    "todos": ["（担当者名）タスク1"],
    "concerns": ["（懸念事項1）"]
  }},
  "flow_narrative_analysis": {{
    "narrative_comment": "（理想的なセールスフローに沿っているかどうかの総評。物語のように解説する）",
    "strength_point": "（例：[関係構築] 相手の成功を祝福し、心理的安全性を確保した点。）",
    "weakness_point": "（例：[価値提案] 顧客の課題解決に繋がらない一方的な商品説明に終始した点。）"
  }},
  "detailed_assessment": {{
    "rapport_building": {{
      "score": "（A〜Dの4段階評価）",
      "comment": "（評価基準に照らした、関係構築フェーズに関する評価コメント）",
      "evidence_quote": "（評価の根拠となった会話のまとまり全体を引用）"
    }},
    "problem_discovery": {{
      "score": "（A〜Dの4段階評価）",
      "comment": "（評価基準に照らした、課題発見フェーズに関する評価コメント）",
      "evidence_quote": "（評価の根拠となった会話のまとまり全体を引用）"
    }},
    "value_addition": {{
      "score": "（A〜Dの4段階評価）",
      "comment": "（評価基準に照らした、価値提案フェーズに関する評価コメント）",
      "evidence_quote": "（評価の根拠となった会話のまとまり全体を引用）"
    }},
    "closing": {{
      "score": "（A〜Dの4段階評価）",
      "comment": "（評価基準に照らした、合意形成とクロージングに関する評価コメント）",
      "evidence_quote": "（評価の根拠となった会話のまとまり全体を引用）"
    }}
  }}
}}
```
"""
    try:
        logging.info("Requesting negotiation analysis from GPT-4o.")
        response = client.chat.completions.create(
            model="gpt-4o",
            response_format={"type": "json_object"},
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            temperature=0.1,
            max_tokens=4090
        )
        logging.info("Successfully received negotiation analysis from GPT-4o.")
        return json.loads(response.choices[0].message.content)
    except Exception as e:
        st.error(f"OpenAI APIでの分析中にエラー: {e}")
        logging.error(f"Error during negotiation analysis: {e}")
        return None


def get_refined_report(original_report, user_instruction):
    """ユーザーの指示に基づき、レポートを修正する"""
    system_prompt = "あなたは優秀なアシスタントです。ユーザーの指示に従って、提供されたレポートを修正してください。必ずレポート全体の構造を維持したまま、指示された箇所のみを修正し、修正後のレポート全文を出力してください。"
    user_prompt = f"""
### 元のレポート:
{original_report}
### 修正指示:
{user_instruction}
"""
    try:
        response = client.chat.completions.create(model="gpt-4o", messages=[{"role": "system", "content": system_prompt}, {"role": "user", "content": user_prompt}])
        return response.choices[0].message.content
    except Exception as e:
        st.error(f"レポートの修正中にエラー: {e}"); return original_report

# -------------------------------------------------------------------
# 4. ヘルパー関数 (Wordファイル生成, DB操作など)
# -------------------------------------------------------------------
def format_timestamp(seconds):
    """秒をHH:MM:SS形式の文字列に変換する"""
    return str(timedelta(seconds=int(seconds)))

def create_minutes_docx(report_text):
    doc = Document()
    doc.add_heading('商談議事録', 0)
    lines = report_text.split('\n')
    for line in lines:
        line = line.strip()
        if line.startswith('### '):
            doc.add_heading(line.replace('### ', ''), level=2)
        elif line.startswith('* **'):
            p = doc.add_paragraph()
            parts = line.replace('* **', '').split('**:', 1)
            run = p.add_run(parts[0])
            run.bold = True
            if len(parts) > 1:
                p.add_run(":" + parts[1])
        elif line.startswith('* '):
            doc.add_paragraph(line.replace('* ', ''), style='List Bullet')
        elif line.strip():
            doc.add_paragraph(line)
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()

def create_analysis_docx(analysis_data, negotiation_info, transcript_display):
    doc = Document()
    doc.add_heading('AI交渉分析レポート', 0)
    
    # 基本情報
    doc.add_paragraph(f"企業名: {negotiation_info.get('client_company', 'N/A')}")
    doc.add_paragraph(f"営業担当: {negotiation_info.get('sales_rep', 'N/A')}")
    doc.add_paragraph(f"日時: {negotiation_info.get('date', 'N/A')}")
    doc.add_paragraph()

    # 総合評価
    score, score_breakdown = calculate_final_score(analysis_data, transcript_display, negotiation_info)
    doc.add_heading(f"総合評価: {score}点", level=1)
    doc.add_paragraph(score_breakdown.replace("\n", " / "))
    
    # 交渉全体の流れ
    narrative = analysis_data.get('flow_narrative_analysis', {})
    doc.add_heading(f"交渉全体の流れ", level=1)
    doc.add_paragraph("評価基準：本レポートでは、交渉を以下の4つのステージに分解し、各ステージの達成度を評価基準としています。\n`関係構築 → 課題発見 → 価値提案 → 合意形成とクロージング`")
    doc.add_paragraph(f"総評: {narrative.get('narrative_comment', '')}")
    
    # 各ステージの詳細評価
    doc.add_heading('交渉の詳細評価', level=1)
    flow = analysis_data.get('detailed_assessment', {})
    stage_map = {
        "rapport_building": "関係構築", "problem_discovery": "課題発見",
        "value_addition": "価値提案", "closing": "合意形成とクロージング"
    }
    for key, stage_name in stage_map.items():
        stage_data = flow.get(key, {})
        if stage_data:
            doc.add_heading(f"{stage_name} (評価: {stage_data.get('score', 'N/A')})", level=2)
            doc.add_paragraph(f"コメント: {stage_data.get('comment', '')}")

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()


def save_report_to_db(negotiation_info, analysis_data, report_markdown, cleaned_transcript):
    """分析結果と最終レポートをSQLiteデータベースに保存する"""
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    c.execute('''
        INSERT INTO reports (timestamp, sales_rep, client_company, client_rep, report_date, analysis_json, report_markdown, cleaned_transcript)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?)
    ''', (
        datetime.now().isoformat(), negotiation_info['sales_rep'], negotiation_info['client_company'],
        negotiation_info['client_rep'], negotiation_info['date'],
        json.dumps(analysis_data, ensure_ascii=False), report_markdown,
        json.dumps(cleaned_transcript, ensure_ascii=False)
    ))
    conn.commit()
    conn.close()
    logging.info(f"Report for {negotiation_info['client_company']} saved to database.")

def calculate_final_score(analysis_json, transcript_display, negotiation_info):
    """AIの質的評価(A-D)と会話バランスから最終スコアを算出する"""
    score_mapping = {"A": 20, "B": 15, "C": 10, "D": 5}
    total_score = 0
    breakdown_texts = []

    assessment = analysis_json.get("detailed_assessment", {})
    stage_map = {
        "rapport_building": "関係構築", "problem_discovery": "課題発見",
        "value_addition": "価値提案", "closing": "合意形成とクロージング"
    }

    for key, name in stage_map.items():
        grade = assessment.get(key, {}).get("score")
        points = score_mapping.get(grade, 0)
        total_score += points
        breakdown_texts.append(f"{name}({grade}評価): {points}点")

    # 会話バランスの計算とスコアリング
    our_company_name = negotiation_info.get('sales_rep', '')
    all_speakers = list(set(item.get('speaker', '') for item in transcript_display))
    our_speaker_label = ''
    our_company_last_name = our_company_name.split(' ')[0][:2]
    for speaker in all_speakers:
        if our_company_last_name in speaker:
            our_speaker_label = speaker
            break
    
    our_company_words = 0
    client_words = 0
    if transcript_display:
        for item in transcript_display:
            word_count = len(re.findall(r'\w+', item.get('text', '')))
            if item.get('speaker', '') == our_speaker_label and our_speaker_label:
                our_company_words += word_count
            else:
                client_words += word_count
    
    total_words = our_company_words + client_words
    balance_points = 0
    our_ratio = 0
    if total_words > 0:
        our_ratio = (our_company_words / total_words) * 100
        ideal_ratio = 25.0
        deviation = abs(our_ratio - ideal_ratio)

        if deviation <= 5:
            balance_points = 20
        elif deviation <= 10:
            balance_points = 15
        elif deviation <= 15:
            balance_points = 10
        elif deviation <= 20:
            balance_points = 5
        else:
            balance_points = 0
        
        total_score += balance_points
        deviation_display = our_ratio - ideal_ratio
        sign = "+" if deviation_display >= 0 else ""
        breakdown_texts.append(f"会話バランス(理想{sign}{deviation_display:.1f}%): {balance_points}点")
    
    score_breakdown = " + ".join(breakdown_texts)
    return total_score, score_breakdown

# -------------------------------------------------------------------
# 5. UI描画: サイドバー
# -------------------------------------------------------------------

with st.sidebar:
    st.header("AI交渉アシスタント")
    st.markdown("---")
    menu_items = {"creation": "商談レポート作成", "history": "過去のレポート", "feedback": "フィードバック"}
    
    for key, value in menu_items.items():
        if st.button(value, use_container_width=True, type="primary" if st.session_state.current_page == key else "secondary"):
            st.session_state.current_page = key
            if 'viewing_report_id' in st.session_state:
                del st.session_state['viewing_report_id']
            st.rerun()

# -------------------------------------------------------------------
# 6. UI描画: メインコンテンツ (ページ切り替え)
# -------------------------------------------------------------------

if st.session_state.current_page == "creation":
    st.title("商談レポート作成")

    if st.session_state.analysis_stage != "initial":
        if st.button("新しいレポートを作成する"):
            st.session_state.confirm_reset = True
    
    if 'confirm_reset' not in st.session_state: st.session_state.confirm_reset = False
    
    if st.session_state.confirm_reset:
        placeholder = st.empty()
        with placeholder.container(border=True):
            st.warning("現在の作業内容は失われます。新しいレポートを作成しますか？")
            col1, col2 = st.columns(2)
            if col1.button("はい、作成する", use_container_width=True, type="primary"):
                reset_creation_page_state()
                st.session_state.confirm_reset = False
                placeholder.empty()
                st.rerun()
            if col2.button("いいえ", use_container_width=True):
                st.session_state.confirm_reset = False
                placeholder.empty()
                st.rerun()

    if st.session_state.analysis_stage == "initial":
        with st.form("upload_form"):
            st.subheader("商談情報の入力とアップロード")
            neg_date = st.date_input("商談日", value=date.today())
            rep_names = ["田中真奈美", "渡辺徹", "小林恭子", "斎藤学", "工藤新一"]
            sales_rep = st.selectbox("営業担当者名", options=rep_names)
            client_company = st.text_input("顧客企業名", placeholder="株式会社デモ")
            client_rep = st.text_input("顧客担当者名", placeholder="商談 花子")
            uploaded_file = st.file_uploader("音声ファイルを選択", type=['wav', 'mp3', 'm4a'])
            submitted = st.form_submit_button("分析を開始する")
            if submitted:
                if not all([sales_rep, client_company, client_rep, uploaded_file]):
                    st.warning("すべての項目を入力し、ファイルをアップロードしてください。")
                else:
                    st.session_state.negotiation_info = {"date": neg_date.strftime('%Y年%m月%d日'), "sales_rep": sales_rep, "client_company": client_company, "client_rep": client_rep}
                    st.session_state.uploaded_file = uploaded_file
                    st.session_state.analysis_stage = 'processing'
                    st.rerun()

    if st.session_state.analysis_stage == 'processing':
        uploaded_file = st.session_state.get('uploaded_file')
        if uploaded_file:
            with st.status("AIアシスタントが分析中です...", expanded=True) as status:
                raw_transcript_text = ""
                try:
                    status.write("ステップ1/4: 音声ファイルを準備中...")
                    audio_bytes = uploaded_file.getvalue()
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".mp3") as tmp:
                        tmp.write(audio_bytes); temp_path = tmp.name
                    
                    audio = AudioSegment.from_file(temp_path)
                    if len(audio) < 3000: # 3秒未満の場合
                        status.update(label="エラー", state="error", expanded=True)
                        st.error("音声ファイルが短すぎます。3秒以上のファイルをアップロードしてください。")
                        st.session_state.analysis_stage = "initial"
                        st.stop()

                    audio = audio.set_frame_rate(16000).set_sample_width(2).set_channels(1)
                    wav_path = temp_path + ".wav"; audio.export(wav_path, format="wav")
                    
                    status.update(label="✅ ステップ1/4: 音声ファイルを準備しました。")
                    status.write("ステップ2/4: 話者を特定中...")
                    diarization_pipeline = Pipeline.from_pretrained("pyannote/speaker-diarization-3.1", use_auth_token=HF_TOKEN)
                    if torch.cuda.is_available(): diarization_pipeline.to(torch.device("cuda"))
                    diarization = diarization_pipeline(wav_path)
                    
                    status.update(label="✅ ステップ2/4: 話者を特定しました。")
                    status.write("ステップ3/4: 文字起こしを実行中...")
                    device = "cuda" if torch.cuda.is_available() else "cpu"
                    whisper_model = whisper.load_model(WHISPER_MODEL, device=device)
                    transcription_result = whisper_model.transcribe(wav_path, word_timestamps=True, language="ja")
                    
                    status.update(label="✅ ステップ3/4: 文字起こしが完了しました。")
                    status.write("ステップ4/4: 文字起こしと話者情報を結合中...")
                    word_timestamps = [word for segment in transcription_result['segments'] for word in segment['words']]
                    if word_timestamps:
                        speaker_turns = [{'start': turn.start, 'end': turn.end, 'speaker': speaker} for turn, _, speaker in diarization.itertracks(yield_label=True)]
                        for word in word_timestamps:
                            word_center = word['start'] + (word['end'] - word['start']) / 2
                            word['speaker'] = next((turn['speaker'] for turn in speaker_turns if turn['start'] <= word_center <= turn['end']), 'UNKNOWN')
                        
                        current_speaker, current_transcript, start_time = word_timestamps[0]['speaker'], "", word_timestamps[0]['start']
                        for word in word_timestamps:
                            if word['speaker'] != current_speaker:
                                raw_transcript_text += f"{current_speaker} ({format_timestamp(start_time)}): {current_transcript.strip()}\n"
                                current_speaker, current_transcript, start_time = word['speaker'], "", word['start']
                            current_transcript += word['word']
                        raw_transcript_text += f"{current_speaker} ({format_timestamp(start_time)}): {current_transcript.strip()}\n"

                    status.update(label="✅ ステップ4/4: 結合が完了しました。")
                    status.write("GPT-4oによる最終分析中...")
                    analysis_result = get_negotiation_analysis(raw_transcript_text, st.session_state.negotiation_info)
                    
                    if analysis_result:
                        status.update(label="分析完了！", state="complete", expanded=False)
                        st.session_state.analysis_data = analysis_result
                        st.session_state.transcript_display = analysis_result.get('cleaned_transcript', [])
                        st.session_state.analysis_stage = 'done'
                        st.session_state.chat_history = [{"role": "assistant", "content": "レポートとAIコーチングを生成しました。"}]
                        st.rerun()
                    else:
                        status.update(label="分析失敗", state="error")
                        st.error("分析に失敗しました。"); st.session_state.analysis_stage = 'initial'

                except Exception as e:
                    status.update(label="エラー発生", state="error")
                    st.error(f"音声処理中にエラーが発生しました: {e}")
                    logging.error(f"Error in audio processing: {e}")
                    st.session_state.analysis_stage = "initial"

                finally:
                    if 'temp_path' in locals() and temp_path and os.path.exists(temp_path): os.remove(temp_path)
                    if 'wav_path' in locals() and wav_path and os.path.exists(wav_path): os.remove(wav_path)

    if st.session_state.analysis_stage == 'done':
        analysis_data = st.session_state.analysis_data
        
        tab1, tab2, tab3 = st.tabs(["📝 議事録レポート", "🤖 AIコーチング", "🗣️ 全文文字起こし"])

        with tab1:
            # (議事録レポートタブのUIは変更なし)
            st.subheader("対話型レポート編集")
            chat_container = st.container(height=200)
            with chat_container:
                for message in st.session_state.chat_history:
                    with st.chat_message(message["role"]):
                        st.markdown(message["content"])
            
            if prompt := st.chat_input("レポートの修正指示を入力"):
                st.session_state.chat_history.append({"role": "user", "content": prompt})
                with st.spinner("AIがレポートを修正中です..."):
                    refined_report = get_refined_report(st.session_state.report_for_display, prompt)
                    st.session_state.report_for_display = refined_report
                st.session_state.chat_history.append({"role": "assistant", "content": "レポートを修正しました。"})
                st.rerun()
            
            st.subheader("生成レポート")
            if not st.session_state.report_for_display:
                report_data = analysis_data.get('summary_report', {})
                overview = report_data.get('overview', {})
                attendees = overview.get('attendees', {})
                summary_items = report_data.get('summary', [])
                summary_text = "\n".join(f"* {item}" for item in summary_items) if isinstance(summary_items, list) else f"* {summary_items}"

                report_parts = [
                    f"### 1. 商談概要", f"* **日時**: {overview.get('date', 'N/A')}", f"* **出席者**:",
                    f"  * **{attendees.get('client_company', '顧客企業')}**: {attendees.get('client_rep', 'N/A')}",
                    f"  * **弊社**: {attendees.get('our_company', 'N/A')}",
                    f"### 2. 本日の目的（アジェンダ）", f"* {report_data.get('agenda', 'N/A')}",
                    f"### 3. 主要な議論の要約", summary_text,
                    f"### 4. 決定事項", "\n".join(f"* {item}" for item in report_data.get('decisions', ['特になし'])),
                    f"### 5. ToDo（ネクストアクション）", "\n".join(f"* {item}" for item in report_data.get('todos', ['特になし'])),
                    f"### 6. 確認事項・懸念点", "\n".join(f"* {item}" for item in report_data.get('concerns', ['特になし'])),
                ]
                st.session_state.report_for_display = "\n\n".join(report_parts)
            
            edited_report = st.text_area("レポート内容を直接編集", st.session_state.report_for_display, height=400, label_visibility="collapsed")
            if edited_report != st.session_state.report_for_display:
                st.session_state.report_for_display = edited_report
                st.rerun()

        with tab2:
            st.subheader("AIによる交渉分析")
            narrative = analysis_data.get('flow_narrative_analysis', {})
            flow = analysis_data.get('detailed_assessment', {})

            final_score, score_breakdown = calculate_final_score(analysis_data, st.session_state.transcript_display, st.session_state.negotiation_info)
            st.metric("総合評価スコア", f"{final_score} 点", delta=score_breakdown)
            st.markdown("---")
            
            st.markdown(f"##### 交渉全体の流れ")
            st.markdown("**評価基準：** 本レポートでは、交渉を以下の4つのステージに分解し、各ステージの達成度を評価基準としています。\n`関係構築 → 課題発見 → 価値提案 → 合意形成とクロージング`")
            st.info(f"**総評:** {narrative.get('narrative_comment', '')}")
            st.success(f"**良かった点**: {narrative.get('strength_point', '')}")
            st.warning(f"**改善すべき点**: {narrative.get('weakness_point', '')}")
            st.markdown("---")

            st.markdown("##### 交渉の詳細評価")
            stage_map = {
                "rapport_building": "関係構築", "problem_discovery": "課題発見",
                "value_addition": "価値提案", "closing": "合意形成とクロージング"
            }
            for key, stage_name in stage_map.items():
                stage_data = flow.get(key, {})
                if stage_data:
                    with st.expander(f"**{stage_name}** (評価: {stage_data.get('score', 'N/A')})"):
                        st.markdown(f"**コメント:** {stage_data.get('comment', '')}")
                        quote = stage_data.get('evidence_quote', '')
                        formatted_quote = quote.replace('\n', '\n\n> ')
                        st.markdown(f"**根拠となった会話:**\n> {formatted_quote}")
            
            st.markdown("---")
            st.markdown("##### 会話バランス")
            st.caption("理想の会話バランスは、営業担当者25%、顧客75％です。")
            our_company_name = st.session_state.negotiation_info.get('sales_rep', '')
            all_speakers = list(set(item.get('speaker', '') for item in st.session_state.transcript_display))
            our_speaker_label = ''
            our_company_last_name = our_company_name.split(' ')[0][:2]
            for speaker in all_speakers:
                if our_company_last_name in speaker:
                    our_speaker_label = speaker
                    break
            
            our_company_words = 0
            client_words = 0
            if st.session_state.transcript_display:
                for item in st.session_state.transcript_display:
                    word_count = len(re.findall(r'\w+', item.get('text', '')))
                    if item.get('speaker', '') == our_speaker_label and our_speaker_label:
                        our_company_words += word_count
                    else:
                        client_words += word_count
            
            total_words = our_company_words + client_words
            if total_words > 0:
                our_ratio = (our_company_words / total_words) * 100
                client_ratio = (client_words / total_words) * 100
                
                fig = go.Figure(data=[go.Pie(labels=['顧客', '営業担当'], values=[client_ratio, our_ratio], hole=.3, marker_colors=['#636EFA', '#EF553B'])])
                fig.update_traces(textinfo='percent+label', textfont_size=14, hovertemplate='<b>%{label}</b>: %{value:.1f}%<extra></extra>')
                fig.update_layout(height=300, margin=dict(t=10, b=10, l=10, r=10), showlegend=False)
                st.plotly_chart(fig, use_container_width=True)

                if 20 <= our_ratio <= 30:
                    st.success("✔️ **理想的な会話バランスです。** 顧客の話を十分に引き出し、効果的な対話ができています。")
                elif our_ratio > 30:
                    st.warning("⚠️ **営業担当者の発話が多めです。** 次回は、質問を増やして顧客が話す時間を確保することを意識しましょう。")
                else:
                    st.warning("⚠️ **顧客の話を引き出す余地があります。** オープンな質問を投げかけ、より積極的に対話をリードしましょう。")

        with tab3:
            # (全文文字起こしタブのUIは変更なし)
            st.subheader("全文文字起こし")
            transcript_container = st.container(height=600)
            with transcript_container:
                for item in st.session_state.transcript_display:
                    st.markdown(f"**{item.get('speaker', '不明')}** ({item.get('start_time', '00:00:00')}): {item.get('text', '')}")

        st.sidebar.markdown("---")
        st.sidebar.subheader("保存とダウンロード")
        
        def save_current_report():
            if not st.session_state.get('report_saved', False):
                save_report_to_db(st.session_state.negotiation_info, st.session_state.analysis_data, st.session_state.report_for_display, st.session_state.transcript_display)
                st.session_state.report_saved = True
                st.toast("レポートを履歴に保存しました！")

        minutes_docx = create_minutes_docx(st.session_state.report_for_display)
        st.sidebar.download_button("議事録ダウンロード", minutes_docx, "議事録.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True, on_click=save_current_report)
        
        analysis_docx = create_analysis_docx(analysis_data, st.session_state.negotiation_info, st.session_state.transcript_display)
        st.sidebar.download_button("AI分析レポートダウンロード", analysis_docx, "AI分析レポート.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True, on_click=save_current_report)


elif st.session_state.current_page == "history":
    # (過去のレポートページのUIは変更なし)
    st.title("過去のレポート一覧")
    
    if 'viewing_report_id' in st.session_state and st.session_state.viewing_report_id is not None:
        report_id = st.session_state.get("viewing_report_id")
        conn = sqlite3.connect(DB_FILE)
        c = conn.cursor()
        c.execute("SELECT report_markdown, analysis_json, cleaned_transcript FROM reports WHERE id = ?", (report_id,))
        data = c.fetchone()
        conn.close()
        
        if data:
            report_markdown, analysis_json_str, cleaned_transcript_str = data
            
            st.subheader("レポート閲覧")
            st.markdown(report_markdown)
            st.markdown("---")
            
            if st.button("このレポートを編集する", type="primary"):
                analysis_data = json.loads(analysis_json_str)
                st.session_state.analysis_data = analysis_data
                st.session_state.report_for_display = report_markdown
                overview = analysis_data.get('summary_report', {}).get('overview', {})
                attendees = overview.get('attendees', {})
                st.session_state.negotiation_info = {
                    "date": overview.get('date', 'N/A'),
                    "sales_rep": attendees.get('our_company', 'N/A'),
                    "client_company": attendees.get('client_company', 'N/A'),
                    "client_rep": attendees.get('client_rep', 'N/A')
                }
                st.session_state.transcript_display = json.loads(cleaned_transcript_str) if cleaned_transcript_str else []
                st.session_state.analysis_stage = "done"
                st.session_state.current_page = "creation"
                st.session_state.report_saved = True
                del st.session_state['viewing_report_id']
                st.rerun()

            if st.button("レポート一覧に戻る"):
                del st.session_state['viewing_report_id']
                st.rerun()

    else:
        conn = sqlite3.connect(DB_FILE)
        c = conn.cursor()
        c.execute("SELECT id, report_date, sales_rep, client_company FROM reports ORDER BY timestamp DESC")
        all_reports = c.fetchall()
        conn.close()

        if not all_reports: st.info("保存されているレポートはありません。")
        
        for report in all_reports:
            report_id, report_date, sales_rep, client_company = report
            with st.container(border=True):
                st.subheader(f"{client_company}様")
                st.write(f"担当: {sales_rep} | 日付: {report_date}")
                if st.button("このレポートを表示する", key=f"open_{report_id}"):
                    st.session_state.viewing_report_id = report_id
                    st.rerun()

elif st.session_state.current_page == "feedback":
    st.title("営業担当者フィードバック")
    rep_names = ["田中真奈美", "渡辺徹", "小林恭子", "斎藤学", "工藤新一"]
    selected_name = st.selectbox("フィードバックを見る担当者を選択してください", options=rep_names)
    
    if selected_name:
        conn = sqlite3.connect(DB_FILE)
        c = conn.cursor()
        c.execute("SELECT analysis_json, report_date, client_company, cleaned_transcript FROM reports WHERE sales_rep = ? ORDER BY timestamp DESC", (selected_name,))
        user_reports_data = c.fetchall()
        conn.close()
        
        if not user_reports_data:
            st.warning(f"{selected_name}さんのレポートは見つかりませんでした。")
        else:
            total_scores = []
            for report in user_reports_data:
                analysis_data = json.loads(report[0])
                transcript_display = json.loads(report[3]) if report[3] else []
                negotiation_info_feedback = {"sales_rep": selected_name}
                score, _ = calculate_final_score(analysis_data, transcript_display, negotiation_info_feedback)
                total_scores.append(score)

            if total_scores:
                avg_score = sum(total_scores) / len(total_scores)
                st.success(f"{len(user_reports_data)}件の商談データに基づき、フィードバックを生成しました。")
                st.metric("平均総合評価スコア", f"{avg_score:.1f} 点")
                
                if avg_score >= 80:
                    st.info("素晴らしい成績です！安定して質の高い交渉ができています。")
                elif avg_score >= 60:
                    st.info("安定した交渉ができています。次は付加価値提案の質を高めることを意識してみましょう。")
                else:
                    st.warning("改善の余地があります。まずは顧客の課題発見に注力し、共感を示すことから始めましょう。")
            else:
                st.warning("有効なスコアデータが見つかりませんでした。")

            st.markdown("---")
            st.subheader("過去のAIコーチングフィードバック一覧")

            for i, report_data in enumerate(user_reports_data):
                analysis_data = json.loads(report_data[0])
                report_date = report_data[1]
                client_company = report_data[2]
                
                narrative = analysis_data.get('flow_narrative_analysis', {})
                
                with st.expander(f"**{report_date}** - **{client_company}様**"):
                    st.markdown(f"**交渉の流れ:** {narrative.get('narrative_comment', 'N/A')}")
